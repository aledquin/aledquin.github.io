#!/depot/Python/Python-3.8.0/bin/python

import os
import pathlib
import sys
import argparse
from io import StringIO

import pytest
from unittest import mock
import docx
import pandas as pd
# ---------------------------------- #
bindir = str(pathlib.Path(__file__).resolve().parent)
sys.path.append(bindir + '/../../bin/')
sys.path.append(bindir + '/../lib/python/Util')
# ---------------------------------- #
import pin_check__ckt_specs
import CommonHeader


@pytest.fixture(autouse=True)
def setup_teardown():
    with mock.patch('sys.stdout', new=StringIO()):
        # logger from Messaging.py is global, but doesn't exist here
        pin_check__ckt_specs.logger = ''

        args = parse_cmd_args([])
        CommonHeader.init(args, [], [])
        yield
        clean_up()


def clean_up():
    if os.path.isfile(os.getcwd() + "/temp_test_file__pin_check.test"):
        os.unlink(os.getcwd() + "/temp_test_file__pin_check.test")
    if os.path.isfile(os.getcwd() + "/temp_pin_info_file__pin_check.csv"):
        os.unlink(os.getcwd() + "/temp_pin_info_file__pin_check.csv")
    if os.path.isfile(os.getcwd() + "/temp_spec_file__pin_check.docx"):
        os.unlink(os.getcwd() + "/temp_spec_file__pin_check.docx")
    if os.path.isfile(os.getcwd() + "/temp_spec_excel__pin_check.xlsx"):
        os.unlink(os.getcwd() + "/temp_spec_excel__pin_check.xlsx")


def parse_cmd_args(tst):
    parser = argparse.ArgumentParser(
        description='Test parser'
    )
    parser.add_argument(
        '-v', metavar='<#>', type=int, default=0, help='verbosity'
    )
    parser.add_argument('-d', metavar='<#>', type=int, default=0, help='debug')   # noqa: E501
    args = parser.parse_args(tst)
    return args


def generate_df(fdocx):
    header_row = 0
    path_spec = f'{bindir}/../data/pincheck/lp5xspecs/{fdocx}'
    document = docx.Document(path_spec)
    for table in document.tables:
        for cell in table.rows[0].cells:
            if cell.text.lower() == 'pin name':
                df = [
                    ['' for i in range(len(table.columns))] for j in range(len(table.rows))   # noqa: E501
                ]
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        if cell.text:
                            df[i][j] = cell.text.replace('\n', '')
                dfs = pd.DataFrame(df)
                dfs.drop_duplicates(subset=None, keep='first', inplace=False)   # noqa: E501
                dfs.columns = dfs.iloc[header_row]
                dfs.columns = dfs.columns.str.upper()
                dfs = dfs.drop(labels=0, axis=0)
                if 'I/O' in dfs.columns:
                    dfs.rename(columns={'I/O': 'DIRECTION'}, inplace=True)   # noqa: E501

                if 'WIDTH' in dfs.columns:
                    dfs.rename(columns={'WIDTH': 'SIGNAL WIDTH'}, inplace=True)   # noqa: E501

                sub1 = dfs[['DIRECTION', 'PIN NAME', 'SIGNAL WIDTH']]

                d_str = ['direction', 'i/o']
                sub1 = sub1[~sub1['DIRECTION'].str.contains('|'.join(d_str), case=False)]   # noqa: E501
                sub1 = sub1[~sub1['PIN NAME'].str.contains('pin name', case=False)]   # noqa: E501
                w_str = ['signal width', 'width']
                sub1 = sub1[
                    ~sub1['SIGNAL WIDTH'].str.contains('|'.join(w_str), case=False)   # noqa: E501
                ]

                sub = sub1.apply(lambda x: x.str.strip() if x.dtype == 'object' else x)   # noqa: E501
    return sub


def generate_pin_info_df(fcsv):
    pin_path = f'{bindir}/../data/pincheck/lp5xspecs/{fcsv}'
    df_pin = pd.read_csv(pin_path, header=0)
    # printing a csv from perforce adds and extra header line
    df_pin.columns = df_pin.columns.str.upper()
    if "I/O" in df_pin.columns:
        df_pin.rename(columns={"I/O": "DIRECTION"}, inplace=True)
    if "PIN NAME" in df_pin.columns:
        df_pin.rename(columns={"PIN NAME": "NAME"}, inplace=True)
    df_pin = df_pin[["NAME", "DIRECTION"]]
    return df_pin


def test_analyze_excel_table():
    header_row = 0
    fdocx = 'dwc_lpddr5xphy_rxacvref_spec#10.docx'
    fcsv = 'dwc_lpddr5xphy_rxacvref_ew#2.csv'
    path_spec = f'{bindir}/../data/pincheck/lp5xspecs/{fdocx}'
    pin_path = f'{bindir}/../data/pincheck/lp5xspecs/{fcsv}'
    document = docx.Document(path_spec)
    for table in document.tables:
        for cell in table.rows[0].cells:
            if cell.text.lower() == 'pin name':
                expected = []
                with mock.patch.object(pin_check__ckt_specs, 'no_stdout'):
                    answer = pin_check__ckt_specs.analyse_excel_table(table, pin_path, header_row)   # noqa: E501
                    assert answer == expected


def test_validate():
    error = ''
    answer = pin_check__ckt_specs.validate(error)
    expected = ''
    assert answer == expected

    error = 'some error'
    answer = pin_check__ckt_specs.validate(error)
    expected = 'some error'
    assert answer == expected


def test_check_width():
    sub = generate_df('dwc_lpddr5xphy_rxacvref_spec#10.docx')   # noqa: E501
    errors = []
    answer = pin_check__ckt_specs.check_width(sub, errors)
    expected = []
    assert answer == expected

    sub = generate_df('dwc_lpddr5xphy_rxacvref_spec-with-errors.docx')   # noqa: E501
    errors = []
    answer = pin_check__ckt_specs.check_width(sub, errors)
    expected = ['Signal Width Error']
    assert answer == expected


def test_check_direction():
    sub = generate_df('dwc_lpddr5xphy_rxacvref_spec#10.docx')   # noqa: E501
    errors = []
    answer = pin_check__ckt_specs.check_direction(sub, errors)
    expected = []
    assert answer == expected

    sub = generate_df('dwc_lpddr5xphy_rxacvref_spec-with-errors.docx')   # noqa: E501
    errors = []
    answer = pin_check__ckt_specs.check_direction(sub, errors)
    expected = ['Direction Error']
    assert answer == expected


def test_read_excel():
    sub = generate_df('dwc_lpddr5xphy_rxacvref_spec#10.docx')   # noqa: E501
    errors = []
    pin_path = f'{bindir}/../data/pincheck/lp5xspecs/dwc_lpddr5xphy_rxacvref_ew#2.csv'   # noqa: E501
    with mock.patch.object(pin_check__ckt_specs, 'no_stdout'):
        answer = pin_check__ckt_specs.read_excel(sub, errors, pin_path)
        expected = []
        assert answer == expected


def test_check_match_width():
    errors = []
    df_spec = generate_df('dwc_lpddr5xphy_rxacvref_spec-with-errors.docx')
    df_pin = generate_pin_info_df('dwc_lpddr5xphy_rxacvref_ew#2.csv')
    df_spec.loc[df_spec['SIGNAL WIDTH'] == 'test', ['SIGNAL WIDTH']] = 2
    with mock.patch.object(pin_check__ckt_specs, 'no_stdout'):
        expected = ['Pin Info Error']
        answer = pin_check__ckt_specs.check_match_width(df_spec, df_pin, errors)
        assert answer == expected


def test_check_match_direction():
    errors = []
    df_spec = generate_df('dwc_lpddr5xphy_rxacvref_spec-with-errors.docx')
    df_pin = generate_pin_info_df('dwc_lpddr5xphy_rxacvref_ew#2.csv')
    df_spec.loc[df_spec['SIGNAL WIDTH'] == 'test', ['DIRECTION']] = 'O'

    expected = ['Direction Mismatch Error']
    answer = pin_check__ckt_specs.check_match_direction(df_spec, df_pin, errors)
    assert answer == expected


def test_check_missing():
    errors = []
    df_spec = generate_df('dwc_lpddr5xphy_rxacvref_spec-with-errors.docx')
    df_pin = generate_pin_info_df('dwc_lpddr5xphy_rxacvref_ew#2.csv')
    df_spec.loc[df_spec['SIGNAL WIDTH'] == 'test', ['PIN NAME']] = 'TEST'

    expected = ['Spec Missing Pin Error', 'Pin Info Missing Pin Error']
    answer = pin_check__ckt_specs.check_missing(df_spec, df_pin, errors)
    assert answer == expected


def test_rename_cols():
    dfs = generate_df('dwc_lpddr5xphy_rxacvref_spec-with-errors.docx')
    dfs.rename(columns={"DIRECTION": "I/O"}, inplace=True)
    dfs.rename(columns={"SIGNAL WIDTH": "WIDTH"}, inplace=True)

    pin_check__ckt_specs.rename_cols(dfs)


def test_verify_args():
    pin_path = f'{bindir}/../data/pincheck/lp5xspecs/dwc_lpddr5xphy_rxacvref_ew#2.csv'   # noqa: E501
    spec_path = f'{bindir}/../data/pincheck/lp5xspecs/dwc_lpddr5xphy_rxacvref_spec#10.docx'   # noqa: E501

    answer = pin_check__ckt_specs.verify_args(pin_path, spec_path)
    expected = (False, False)
    assert answer == expected

    pin_path = '//fakep4path/test.csv'
    spec_path = '//fakep4path/test.docx'

    answer = pin_check__ckt_specs.verify_args(pin_path, spec_path)
    expected = (True, True)
    assert answer == expected


def test_verify_p4_path():
    with mock.patch("P4.P4.disconnect") as mock_disconnect:
        with pytest.raises(SystemExit):
            file_path = '//depot/fake/path'
            pin_check__ckt_specs.verify_p4_path(file_path)
        pin_check__ckt_specs.verify_p4_path('//depot')
        mock_disconnect.assert_called()


def test_verify_local_path():
    with pytest.raises(SystemExit):
        file_path = '/local/fake/path'
        pin_check__ckt_specs.verify_local_path(file_path)


def test_macro_name():
    pin_path = f'{bindir}/../data/pincheck/lp5xspecs/dwc_lpddr5xphy_rxacvref_ew#2.csv'   # noqa: E501
    spec_path = f'{bindir}/../data/pincheck/lp5xspecs/dwc_lpddr5xphy_rxacvref_spec#10.docx'   # noqa: E501
    answer = pin_check__ckt_specs.macro_name(pin_path, spec_path)
    expected = ('rxacvref', 0)
    assert answer == expected

    with mock.patch.object(pin_check__ckt_specs, 'isolate_macro', return_value=[]) as mock_isolate:   # noqa: E501
        with pytest.raises(SystemExit):
            pin_path = 'dwc_lpddr5xphy_rxacvref_ew#2.csv'
            spec_path = 'dwc_lpddr5xphy_rxacvref_spec#10.docx'
            pin_check__ckt_specs.macro_name(pin_path, spec_path)
        mock_isolate.assert_called()


def test_match_macro():
    pin_macros = ['test']
    spec_macros = ['test']

    answer = pin_check__ckt_specs.match_macro(pin_macros, spec_macros)
    expected = 'test'
    assert answer == expected

    with pytest.raises(SystemExit):
        spec_macros = ['']
        pin_check__ckt_specs.match_macro(pin_macros, spec_macros)


def test_isolate_macro():
    array = ['ew', 'spec', 'test']

    answer = pin_check__ckt_specs.isolate_macro(array)
    expected = ['test']
    assert answer == expected


def test_header_macro():
    macro = 'test'
    tex = 'test'
    expected = True

    answer = pin_check__ckt_specs.header_macro(macro, tex)
    assert answer == expected

    tex = 'dwc_lpddr5xphy_test'
    answer = pin_check__ckt_specs.header_macro(macro, tex)
    assert answer == expected


def test_p4_file():
    with mock.patch("Misc.run_system_cmd") as mock_cmd:
        mock_cmd.return_value = 'new/path'
        perforce_path = 'test'
        name = 'test'
        ext = 'test'
        with mock.patch('os.path.abspath', return_value='new/path'):  # noqa: E501
            answer = pin_check__ckt_specs.p4_file(perforce_path, name, ext)
            expected = 'new/path'
            assert answer == expected


def test_analyze_document():
    header_row = 1
    macro = 'lstx_dx4'
    fdocx = 'dwc_lpddr5xphy_lstxacx2_lstxcsx2_lstxdx4_lstxdx5_lstxzcal_spec_copy.docx'   # noqa: E501
    fcsv = 'dwc_lpddr5xphy_lstx_dx4_ew.csv'
    spec_path = f'{bindir}/../data/pincheck/lp5xspecs/{fdocx}'
    pin_path = f'{bindir}/../data/pincheck/lp5xspecs/{fcsv}'
    document = docx.Document(spec_path)
    with mock.patch.object(pin_check__ckt_specs, 'analyse_excel_table') as mock_analyze_excel:   # noqa: E501
        pin_check__ckt_specs.analyze_document(document, macro, header_row, pin_path)   # noqa: E501
        mock_analyze_excel.assert_called()

    macro = 'not a macro'
    with mock.patch.object(pin_check__ckt_specs, 'eprint') as mock_eprint:   # noqa: E501
        pin_check__ckt_specs.analyze_document(document, macro, header_row, pin_path)   # noqa: E501
        mock_eprint.assert_called()


def test_rm_temp_files():
    with mock.patch('os.path.isfile', return_value=True):
        with mock.patch('os.unlink') as mock_unlink:
            pin_check__ckt_specs.rm_temp_files()
            mock_unlink.assert_called()


if __name__ == "__main__":
    # nolint utils__script_usage_statistics
    # nolint main
    ret_code = pytest.main([__file__] + sys.argv[1:])
    sys.exit(ret_code)
